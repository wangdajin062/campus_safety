"""
latex2docx.py — LaTeX to DOCX converter for paper1_en_v6.tex
Preserves formatting: sections, bold/italic, tables, equations, citations.
Output: docs/ch.v6.docx
"""
import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PROJECT_ROOT = Path(__file__).resolve().parent.parent
TEX_FILE = PROJECT_ROOT / "docs" / "paper1_en_v6.tex"
OUTPUT_FILE = PROJECT_ROOT / "docs" / "ch.v6.docx"

# ── LaTeX command stripping ──────────────────────────────
def strip_latex(text: str) -> str:
    """Remove LaTeX commands and return plain-ish text."""
    # Remove comments
    text = re.sub(r'(?<!\\)%.*$', '', text, flags=re.MULTILINE)
    # Replace common LaTeX constructs
    text = text.replace(r'\textbf{', '\x01').replace(r'\textit{', '\x02')
    text = text.replace(r'\emph{', '\x02')
    text = text.replace(r'\texttt{', '\x03')
    text = text.replace(r'\textsuperscript{', '\x04')
    text = text.replace(r'\textsubscript{', '\x05')
    # Remove \cite{...}
    text = re.sub(r'~?\\cite\{([^}]+)\}', r' [\1]', text)
    # Remove \ref{...}
    text = re.sub(r'~?\\ref\{([^}]+)\}', r' \1', text)
    # Remove \label{...}
    text = re.sub(r'\\label\{[^}]+\}', '', text)
    # Handle \S (section symbol) — keep as §
    text = text.replace(r'\S', '§')
    # Remove \footnote{...}
    text = re.sub(r'\\footnote\{([^}]*(?:\{[^}]*\}[^}]*)*)\}', r' [Footnote: \1]', text)
    # Replace inline math $...$ with formatted text
    text = re.sub(r'\$([^$]+)\$', r'\1', text)
    # Replace displayed math \[...\]
    text = re.sub(r'\\\[.*?\\\]', '[Equation]', text, flags=re.DOTALL)
    # Remove \begin{equation}...\end{equation}
    text = re.sub(r'\\begin\{equation\}.*?\\end\{equation\}', '[Equation]', text, flags=re.DOTALL)
    # Handle \noindent
    text = text.replace(r'\noindent', '')
    # Handle \medskip, \bigskip
    text = re.sub(r'\\(?:med|big)skip', '', text)
    # Handle ~ (non-breaking space)
    text = text.replace('~', ' ')
    # Handle \&
    text = text.replace(r'\&', '&')
    # Handle \%
    text = text.replace(r'\%', '%')
    # Handle \_
    text = text.replace(r'\_', '_')
    # Handle \{ \}
    text = text.replace(r'\{', '{').replace(r'\}', '}')
    # Handle \#
    text = text.replace(r'\#', '#')
    # Handle \textellipsis or \dots
    text = text.replace(r'\dots', '...')
    text = re.sub(r'\\ldots', '...', text)
    # Handle \times
    text = text.replace(r'\times', '×')
    # Handle accented chars
    text = text.replace(r'\"{o}', 'ö').replace(r'\"{u}', 'ü')
    text = text.replace(r'\'{e}', 'é').replace(r'\`{e}', 'è')
    # Remove remaining single-char commands like \,
    text = re.sub(r'\\,', '', text)
    text = re.sub(r'\\;', ' ', text)
    # Remove \textbf, \textit markers we added
    text = text.replace('\x01', '').replace('\x02', '').replace('\x03', '')
    text = text.replace('\x04', '').replace('\x05', '')
    # Collapse multiple spaces
    text = re.sub(r' +', ' ', text)
    # Collapse multiple newlines
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


def add_formatted_paragraph(doc, text: str, style: str = 'Normal', bold: bool = False):
    """Add a paragraph with basic LaTeX inline formatting preserved."""
    para = doc.add_paragraph(style=style)

    # Split by inline formatting markers
    parts = re.split(r'(\\textbf\{|\\textit\{|\\emph\{|\\texttt\{)', text)

    formatting_stack = []
    i = 0
    while i < len(parts):
        part = parts[i]
        if part == r'\textbf{':
            formatting_stack.append('bold')
            i += 1
            continue
        elif part == r'\textit{' or part == r'\emph{':
            formatting_stack.append('italic')
            i += 1
            continue
        elif part == r'\texttt{':
            formatting_stack.append('mono')
            i += 1
            continue

        # Count closing braces in this part
        brace_count = part.count('}')
        # Remove closing braces at end
        clean_text = part

        run = para.add_run(strip_latex(clean_text))

        # Apply current formatting
        fmt = formatting_stack[-1] if formatting_stack else None
        if fmt == 'bold':
            run.bold = True
        elif fmt == 'italic':
            run.italic = True
        elif fmt == 'mono':
            run.font.name = 'Consolas'

        # Pop formatting for each }
        for _ in range(brace_count):
            if formatting_stack:
                formatting_stack.pop()

        i += 1

    return para


def add_section_heading(doc, text: str, level: int):
    """Add a section heading."""
    text = strip_latex(text)
    heading = doc.add_heading(text, level=level)
    return heading


def convert_latex_to_docx(tex_path: Path, output_path: Path):
    """Main conversion function."""
    doc = Document()

    # ── Page setup ──
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    # ── Style setup ──
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(3)

    for i in range(1, 4):
        heading_style = doc.styles[f'Heading {i}']
        heading_style.font.name = 'Times New Roman'
        heading_style.font.color.rgb = RGBColor(0, 0, 0)
        if i == 1:
            heading_style.font.size = Pt(16)
        elif i == 2:
            heading_style.font.size = Pt(13)
        else:
            heading_style.font.size = Pt(11.5)

    # ── Read LaTeX source ──
    with open(tex_path, 'r', encoding='utf-8') as f:
        content = f.read()

    # Remove preamble (everything before \begin{document})
    doc_start = content.find(r'\begin{document}')
    if doc_start >= 0:
        content = content[doc_start + len(r'\begin{document}'):]

    # Remove \end{document}
    doc_end = content.find(r'\end{document}')
    if doc_end >= 0:
        content = content[:doc_end]

    # ── Title ──
    title_match = re.search(r'\\title\{((?:[^{}]|\{[^{}]*\})*)\}', content)
    if title_match:
        title_text = strip_latex(title_match.group(1))
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(title_text)
        title_run.bold = True
        title_run.font.size = Pt(18)

    # ── Author ──
    author_match = re.search(r'\\author\{([^}]+)\}', content)
    if author_match:
        author_text = strip_latex(author_match.group(1))
        author_para = doc.add_paragraph()
        author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        author_para.add_run(author_text).font.size = Pt(11)

    # ── Abstract ──
    abstract_match = re.search(r'\\begin\{abstract\}(.*?)\\end\{abstract\}', content, re.DOTALL)
    if abstract_match:
        doc.add_heading('Abstract', level=2)
        abstract_text = strip_latex(abstract_match.group(1))
        for para_text in abstract_text.split('\n\n'):
            if para_text.strip():
                p = doc.add_paragraph()
                p.add_run(para_text.strip()).font.size = Pt(10)
                p.paragraph_format.left_indent = Cm(1)
                p.paragraph_format.right_indent = Cm(1)

    # ── Keywords ──
    kw_match = re.search(r'\\begin\{keyword\}(.*?)\\end\{keyword\}', content, re.DOTALL)
    if kw_match:
        kw_text = strip_latex(kw_match.group(1))
        p = doc.add_paragraph()
        run = p.add_run('Keywords: ')
        run.bold = True
        run.font.size = Pt(10)
        p.add_run(kw_text).font.size = Pt(10)

    doc.add_page_break()

    # ── Process body content ──
    # Remove frontmatter
    fm_end = content.find(r'\end{frontmatter}')
    if fm_end >= 0:
        body = content[fm_end + len(r'\end{frontmatter}'):]
    else:
        body = content

    # Remove appendix content (process separately)
    app_start = body.find(r'\appendix')
    if app_start >= 0:
        appendix_body = body[app_start + len(r'\appendix'):]
        body = body[:app_start]
    else:
        appendix_body = ""

    # Process body line by line
    _process_body(doc, body)

    # ── Appendix ──
    if appendix_body.strip():
        doc.add_page_break()
        doc.add_heading('Appendix', level=1)
        _process_body(doc, appendix_body)

    # ── Save ──
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"DOCX saved: {output_path}")
    return output_path


def _process_body(doc, body: str):
    """Process LaTeX body content into DOCX paragraphs."""
    lines = body.split('\n')

    i = 0
    in_table = False
    in_figure = False
    in_equation = False
    in_list = False
    table_lines = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip empty lines
        if not stripped:
            if in_table or in_figure or in_equation:
                i += 1
                continue
            i += 1
            continue

        # Section headings
        sec_match = re.match(r'\\section\{((?:[^{}]|\{[^{}]*\})*)\}', stripped)
        if sec_match:
            add_section_heading(doc, sec_match.group(1), 1)
            i += 1
            continue

        sub_match = re.match(r'\\subsection\{((?:[^{}]|\{[^{}]*\})*)\}', stripped)
        if sub_match:
            add_section_heading(doc, sub_match.group(1), 2)
            i += 1
            continue

        subsub_match = re.match(r'\\subsubsection\{((?:[^{}]|\{[^{}]*\})*)\}', stripped)
        if subsub_match:
            add_section_heading(doc, subsub_match.group(1), 3)
            i += 1
            continue

        # Table start
        if re.match(r'\\begin\{table', stripped) or re.match(r'\\begin\{table\*', stripped):
            in_table = True
            table_lines = []
            i += 1
            continue

        # Figure start
        if re.match(r'\\begin\{figure', stripped) or re.match(r'\\begin\{figure\*', stripped):
            in_figure = True
            i += 1
            continue

        # Equation display start
        if re.match(r'\\begin\{equation', stripped) or stripped == r'\[':
            in_equation = True
            i += 1
            continue

        # Itemize/enumerate start
        if re.match(r'\\begin\{(?:itemize|enumerate)\}', stripped):
            in_list = True
            i += 1
            continue

        # End environments
        if re.match(r'\\end\{table', stripped) or re.match(r'\\end\{table\*', stripped):
            in_table = False
            if table_lines:
                _process_table(doc, '\n'.join(table_lines))
                table_lines = []
            i += 1
            continue

        if re.match(r'\\end\{figure', stripped) or re.match(r'\\end\{figure\*', stripped):
            in_figure = False
            i += 1
            continue

        if re.match(r'\\end\{equation', stripped) or stripped == r'\]':
            in_equation = False
            i += 1
            continue

        if re.match(r'\\end\{(?:itemize|enumerate)\}', stripped):
            in_list = False
            i += 1
            continue

        # Collect table content
        if in_table:
            table_lines.append(line)
            i += 1
            continue

        # Skip figure content
        if in_figure:
            i += 1
            continue

        # Skip equation content
        if in_equation:
            i += 1
            continue

        # List item
        if in_list and stripped.startswith(r'\item'):
            item_text = re.sub(r'^\\item\s*', '', stripped)
            item_text = strip_latex(item_text)
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(item_text)
            i += 1
            continue

        # \caption
        cap_match = re.match(r'\\caption\{((?:[^{}]|\{[^{}]*\})*)\}', stripped)
        if cap_match:
            cap_text = strip_latex(cap_match.group(1))
            p = doc.add_paragraph()
            run = p.add_run(cap_text)
            run.italic = True
            run.font.size = Pt(10)
            i += 1
            continue

        # \includegraphics → placeholder
        if stripped.startswith(r'\includegraphics'):
            p = doc.add_paragraph()
            run = p.add_run('[Figure]')
            run.italic = True
            run.font.color.rgb = RGBColor(128, 128, 128)
            i += 1
            continue

        # Bold paragraph headings like \noindent\textbf{...}
        bold_heading = re.match(r'\\noindent\\textbf\{((?:[^{}]|\{[^{}]*\})*)\}', stripped)
        if bold_heading:
            p = doc.add_paragraph()
            run = p.add_run(strip_latex(bold_heading.group(1)))
            run.bold = True
            i += 1
            continue

        # Standalone \textbf{...} paragraph
        standalone_bold = re.match(r'\\textbf\{((?:[^{}]|\{[^{}]*\})*)\}', stripped)
        if standalone_bold:
            p = doc.add_paragraph()
            run = p.add_run(strip_latex(standalone_bold.group(1)))
            run.bold = True
            i += 1
            continue

        # \item with \textbf (numbered items)
        item_bold = re.match(r'\\item\s*\\textbf\{((?:[^{}]|\{[^{}]*\})*)\}', stripped)
        if item_bold:
            item_text = strip_latex(item_bold.group(1))
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(item_text)
            run.bold = True
            i += 1
            continue

        # \medskip, \bigskip → spacing
        if re.match(r'\\(?:med|big)skip', stripped):
            doc.add_paragraph()
            i += 1
            continue

        # Regular text paragraph
        text = strip_latex(stripped)
        if text and len(text) > 3:
            add_formatted_paragraph(doc, stripped)

        i += 1


def _process_table(doc, table_content: str):
    """Convert LaTeX table to DOCX table."""
    # Try to extract caption
    cap_match = re.search(r'\\caption\{((?:[^{}]|\{[^{}]*\})*)\}', table_content)
    if cap_match:
        cap_text = strip_latex(cap_match.group(1))
        p = doc.add_paragraph()
        run = p.add_run(f'Table: {cap_text}')
        run.italic = True
        run.font.size = Pt(10)

    # Extract tabular content
    tab_match = re.search(r'\\begin\{tabular\*?\}.*?\n(.*?)\\end\{tabular\*?\}', table_content, re.DOTALL)
    if not tab_match:
        tab_match = re.search(r'\\begin\{tabularx\}.*?\n(.*?)\\end\{tabularx\}', table_content, re.DOTALL)

    if tab_match:
        tab_body = tab_match.group(1)
        rows = []
        for line in tab_body.split('\n'):
            line = line.strip()
            if not line or line.startswith('%'):
                continue
            if '\\toprule' in line or '\\midrule' in line or '\\bottomrule' in line:
                continue
            if line.startswith('\\hline'):
                continue
            if line.startswith('\\rowcolor'):
                line = re.sub(r'\\rowcolor\[[^\]]*\]\{[^}]*\}', '', line).strip()
            if line.startswith(r'\multicolumn'):
                continue

            # Split by & but not inside braces
            cells = []
            depth = 0
            current = ''
            for ch in line:
                if ch == '{': depth += 1
                elif ch == '}': depth -= 1
                elif ch == '&' and depth == 0:
                    cells.append(current.strip())
                    current = ''
                    continue
                elif ch == '\\' and depth == 0:
                    # End of row marker
                    break
                current += ch
            if current.strip():
                cells.append(current.strip())

            if cells:
                rows.append([strip_latex(c) for c in cells])

        if rows:
            table = doc.add_table(rows=len(rows), cols=len(rows[0]))
            table.style = 'Light Grid Accent 1'
            for r_idx, row_data in enumerate(rows):
                for c_idx, cell_text in enumerate(row_data):
                    if c_idx < len(table.rows[r_idx].cells):
                        cell = table.rows[r_idx].cells[c_idx]
                        cell.text = cell_text
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.font.size = Pt(8)

    doc.add_paragraph()  # spacing after table


def main():
    if not TEX_FILE.exists():
        print(f"ERROR: {TEX_FILE} not found")
        return 1

    print(f"Converting: {TEX_FILE.name} → {OUTPUT_FILE.name}")
    convert_latex_to_docx(TEX_FILE, OUTPUT_FILE)
    return 0


if __name__ == "__main__":
    sys.exit(main())
